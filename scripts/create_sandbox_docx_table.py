"""Create or reset the sandbox OneDrive docx with a valid Steven-style order table.

Usage (from repo root, via WSL):
  python scripts/create_sandbox_docx_table.py

Requires .env with:
  MS_CLIENT_ID, MS_TENANT_ID, MS_REFRESH_TOKEN
  ONEDRIVE_TEST_DRIVE_ID, ONEDRIVE_TEST_FILE_ID
  ONEDRIVE_SANDBOX_WRITE_ENABLED=true

Safety:
  - Only writes to ONEDRIVE_TEST_DRIVE_ID / ONEDRIVE_TEST_FILE_ID
  - Refuses if test IDs match ONEDRIVE_DRIVE_ID / ONEDRIVE_FILE_ID
  - Requires ONEDRIVE_SANDBOX_WRITE_ENABLED=true
  - Sandbox file name must include TEST, SANDBOX, COPY, or CLONE
"""
from __future__ import annotations

import io
import os
import sys
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parents[1]
if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from scripts.test_onedrive_sandbox import (
    SafetyRefusal,
    merged_environment,
    require_sandbox_ids,
    require_write_flag,
    apply_runtime_environment,
    client_factory,
    require_sandbox_metadata,
)

COLUMNS = [
    "Date",
    "Item No.",
    "QTY",
    "Color",
    "Customer Name",
    "Sent to Supplier",
    "Ship by date",
    "Sent to customer",
]


def _make_header_cell(cell, text: str) -> None:
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)


def build_sandbox_docx() -> bytes:
    doc = Document()

    # Remove default empty paragraph so document starts with the table
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"

    header_row = table.rows[0]
    for i, col_name in enumerate(COLUMNS):
        _make_header_cell(header_row.cells[i], col_name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def main() -> int:
    env = merged_environment(REPO_ROOT / ".env")

    try:
        require_write_flag(env)
        drive_id, file_id = require_sandbox_ids(env)
        apply_runtime_environment(env)
        client = client_factory(drive_id, file_id)
        metadata = client.get_metadata()
        safe_name = require_sandbox_metadata(metadata)
    except SafetyRefusal as exc:
        print(f"REFUSED: {exc}", file=sys.stderr)
        return 1
    except Exception as exc:
        print(f"FAILED (auth/metadata): {exc}", file=sys.stderr)
        return 1

    print(f"Target sandbox file: {safe_name}")
    print(f"Building fresh docx with {len(COLUMNS)}-column table...")

    try:
        docx_bytes = build_sandbox_docx()
        client.upload_docx(docx_bytes)
    except Exception as exc:
        print(f"FAILED (upload): {exc}", file=sys.stderr)
        return 1

    print("Upload succeeded.")
    print(f"Columns: {', '.join(COLUMNS)}")
    print("The sandbox document now contains an empty Steven-style order table.")
    print("You can now run scripts/test_onedrive_sandbox.py --write-test-row to append a test row.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
