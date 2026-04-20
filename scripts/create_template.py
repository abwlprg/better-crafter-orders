"""Crea el template docxtpl basado en el Word real del cliente."""

from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# 8 columnas visibles (sin la columna vacía final)
HEADER_COLUMNS = [
    "Date", "Item No.", "QTY", "Color",
    "Customer Name", "Sent to Supplier", "Ship by date", "Sent to customer",
]

PLACEHOLDERS = [
    "{{ order_date }}", "{{ item_code }}", "{{ quantity }}", "{{ color }}",
    "{{ customer_name }}", "y", "{{ ship_by }}", "",
]

COL_WIDTHS = [
    Inches(0.9),   # Date
    Inches(1.2),   # Item No.
    Inches(0.6),   # QTY
    Inches(1.8),   # Color
    Inches(2.0),   # Customer Name
    Inches(1.0),   # Sent to Supplier
    Inches(1.0),   # Ship by date
    Inches(1.0),   # Sent to customer
]


def build_template(output_path: Path) -> None:
    """Build a clean 8-column docxtpl template in landscape."""
    doc = Document()

    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    # Swap width/height for landscape
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    # Tighter margins
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # Create table
    num_cols = len(HEADER_COLUMNS)
    table = doc.add_table(rows=2, cols=num_cols)
    table.style = "Table Grid"

    # Set column widths
    for i, width in enumerate(COL_WIDTHS):
        table.columns[i].width = width

    # Header row — bold + centered
    header_row = table.rows[0]
    for i, text in enumerate(HEADER_COLUMNS):
        cell = header_row.cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(11)

    # Template/placeholder row
    data_row = table.rows[1]
    for i, text in enumerate(PLACEHOLDERS):
        cell = data_row.cells[i]
        cell.text = text
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(11)

    # Verify structure
    for ri, row in enumerate(table.rows):
        tcs = row._tr.findall(qn("w:tc"))
        print(f"  Row {ri}: {len(tcs)} tc elements")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✅ Template creado en: {output_path}")


if __name__ == "__main__":
    build_template(Path("templates/stephen_template.docx"))
