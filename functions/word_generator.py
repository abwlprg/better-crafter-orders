"""Word report generation and Firebase Storage upload helpers."""

from __future__ import annotations

import logging
import tempfile
import time
from dataclasses import dataclass
from datetime import date, datetime, timedelta, timezone
from pathlib import Path
from typing import Any

from docx import Document
from docxtpl import DocxTemplate
from firebase_admin import storage

try:
    from . import config
except ImportError:  # Firebase Functions imports modules from functions/ directly.
    import config

logger = logging.getLogger(__name__)


@dataclass(slots=True)
class GeneratedReport:
    """Metadata for an uploaded supplier report."""

    local_path: Path
    storage_path: str
    signed_url: str


class WordReportGenerator:
    """Generates daily supplier order reports from docx templates."""

    def __init__(self, template_path: str) -> None:
        """Create generator with a specific template location."""
        self._template_path = Path(template_path)

    def generate_daily_report(self, orders: list[dict[str, str]], report_date: date) -> Path:
        """Generate a single daily report document containing all order rows."""
        if not orders:
            raise ValueError("At least one order is required to generate a report")

        if not self._template_path.exists():
            raise FileNotFoundError(f"Template file not found: {self._template_path}")

        # Sort orders ascending by date (oldest first → newest at bottom)
        sorted_orders = self._sort_orders_ascending(orders)

        first_order = sorted_orders[0]
        doc_template = DocxTemplate(str(self._template_path))
        context = {
            "order_date":    first_order.get("order_date", ""),
            "item_code":     first_order.get("item_code", ""),
            "quantity":      first_order.get("quantity", ""),
            "color":         first_order.get("color", ""),
            "customer_name": first_order.get("customer_name", ""),
            "ship_by":       first_order.get("ship_by", ""),
            "generated_date": report_date.isoformat(),
        }
        doc_template.render(context)

        temp_file = tempfile.NamedTemporaryFile(
            suffix=f"_{report_date.isoformat()}.docx", delete=False
        )
        temp_file_path = Path(temp_file.name)
        temp_file.close()
        doc_template.save(str(temp_file_path))

        if len(sorted_orders) > 1:
            self._append_orders_to_table(temp_file_path, sorted_orders[1:])

        return temp_file_path

    def upload_report(self, report_path: Path, report_date: date) -> GeneratedReport:
        """Upload report to Firebase Storage and return signed URL metadata."""
        storage_path = self._build_storage_path(report_date)
        bucket = storage.bucket(config.STORAGE_BUCKET)
        blob = bucket.blob(storage_path)

        try:
            blob.upload_from_filename(
                filename=str(report_path),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as error:
            logger.error("Storage upload failed, retrying once", exc_info=error)
            time.sleep(5)
            blob.upload_from_filename(
                filename=str(report_path),
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        signed_url = blob.generate_signed_url(
            version="v4",
            method="GET",
            expiration=timedelta(days=7),
            service_account_email=None,
            access_token=None,
            generation=None,
            response_disposition=None,
            response_type=None,
            query_parameters=None,
        )

        return GeneratedReport(
            local_path=report_path,
            storage_path=storage_path,
            signed_url=signed_url,
        )

    @staticmethod
    def _build_storage_path(report_date: date) -> str:
        """Build destination storage path for the generated report."""
        year = report_date.strftime("%Y")
        month = report_date.strftime("%m")
        day = report_date.strftime("%d")
        file_name = f"stephen_orders_{report_date.isoformat()}.docx"
        return f"{config.REPORT_PREFIX}/{year}/{month}/{day}/{file_name}"

    @staticmethod
    def _sort_orders_ascending(orders: list[dict[str, str]]) -> list[dict[str, str]]:
        """Sort orders by order_date ascending (oldest first, newest at bottom)."""
        def date_key(order: dict[str, str]) -> tuple[int, int]:
            date_str = order.get("order_date", "").strip()
            try:
                parts = date_str.split("/")
                if len(parts) >= 2:
                    month = int(parts[0])
                    day   = int(parts[1])
                    if 1 <= month <= 12 and 1 <= day <= 31:
                        return (month, day)
            except (ValueError, IndexError):
                pass
            return (99, 99)  # malformed dates go to end

        return sorted(orders, key=date_key)

    @staticmethod
    def _append_orders_to_table(file_path: Path, orders: list[dict[str, str]]) -> None:
        """Append additional order rows to the first table in the report."""
        document = Document(str(file_path))
        if not document.tables:
            raise ValueError("Template must contain at least one table")

        table = document.tables[0]
        for order in orders:
            row_cells = table.add_row().cells
            row_cells[0].text = order.get("order_date", "")     # Date
            row_cells[1].text = order.get("item_code", "")      # Item No.
            row_cells[2].text = order.get("quantity", "")       # QTY
            row_cells[3].text = order.get("color", "")          # Color
            row_cells[4].text = order.get("customer_name", "")  # Customer Name
            row_cells[5].text = "y"                              # Sent to Supplier
            row_cells[6].text = order.get("ship_by", "")        # Ship by date
            if len(row_cells) > 7:
                row_cells[7].text = ""                           # Sent to customer

        document.save(str(file_path))


def append_orders_to_existing_docx(docx_bytes: bytes, orders: list[dict[str, str]]) -> tuple[bytes, int, int]:
    """
    Append order rows to the first table in an existing .docx file.

    BUG 2 FIX: Deduplication is now done strictly by Gmail message ID in
    Firestore (see ProcessedEmailStore), NOT by (date, item_code, customer_name).
    The old key collapsed two genuinely distinct orders for the same customer
    (e.g. Bill ordering the same item twice) into a single row. We now trust
    the caller (main.py) to only pass through orders whose Gmail message_id
    has not been processed before, and we append every one of them.

    Args:
        docx_bytes: Raw bytes of the existing Word document (downloaded from OneDrive).
        orders:     List of order dicts with keys order_date, item_code, quantity,
                    color, customer_name, ship_by.

    Returns:
        Tuple of (updated_bytes, appended_count, skipped_count).
        skipped_count is kept for API compatibility and is always 0.
    """
    if not orders:
        raise ValueError("At least one order is required")

    import io

    doc = Document(io.BytesIO(docx_bytes))

    if not doc.tables:
        raise ValueError("The OneDrive document must contain at least one table")

    table = doc.tables[0]

    # Sort ascending so oldest orders come first
    sorted_orders = WordReportGenerator._sort_orders_ascending(orders)

    appended = 0
    skipped = 0
    for order in sorted_orders:
        row_cells = table.add_row().cells
        row_cells[0].text = order.get("order_date", "")      # Date
        row_cells[1].text = order.get("item_code", "")       # Item No.
        row_cells[2].text = order.get("quantity", "")        # QTY
        row_cells[3].text = order.get("color", "")           # Color
        row_cells[4].text = order.get("customer_name", "")   # Customer Name
        row_cells[5].text = "y"                              # Sent to Supplier
        row_cells[6].text = order.get("ship_by", "")         # Ship by date
        if len(row_cells) > 7:
            row_cells[7].text = ""                           # Sent to customer
        logger.info(
            "📝 Appended row: date=%r item=%r qty=%r customer=%r",
            order.get("order_date"), order.get("item_code"),
            order.get("quantity"), order.get("customer_name"),
        )
        appended += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("OneDrive docx: %d appended, %d skipped", appended, skipped)
    return buf.read(), appended, skipped


def clear_rows_from_docx(docx_bytes: bytes, start_date: str = None, end_date: str = None) -> tuple[bytes, int]:
    """
    Delete data rows from the first table of a .docx whose order_date falls
    within [start_date, end_date] (inclusive, format MM/DD/YYYY or MM/DD).
    The header row (row 0) is always preserved.

    If neither date is given, ALL data rows are deleted.

    Returns:
        Tuple of (updated_bytes, deleted_count).
    """
    import io
    from lxml import etree

    doc = Document(io.BytesIO(docx_bytes))
    if not doc.tables:
        raise ValueError("The OneDrive document must contain at least one table")

    table = doc.tables[0]
    deleted = 0

    def _parse_date(val: str):
        """Parse MM/DD/YYYY or MM/DD into (month, day, year). Year defaults to today's year."""
        val = val.strip()
        parts = val.split("/")
        try:
            if len(parts) == 3:
                return (int(parts[2]), int(parts[0]), int(parts[1]))  # (year, month, day)
            elif len(parts) == 2:
                return (datetime.today().year, int(parts[0]), int(parts[1]))
        except (ValueError, IndexError):
            pass
        return None

    def _in_range(order_date: str) -> bool:
        if not start_date and not end_date:
            return True
        d = _parse_date(order_date)
        if d is None:
            return False
        if start_date:
            s = _parse_date(start_date)
            if s and d < s:
                return False
        if end_date:
            e = _parse_date(end_date)
            if e and d > e:
                return False
        return True

    # Iterate rows in reverse (skip row 0 = header) to safely delete
    rows = table.rows
    for i in range(len(rows) - 1, 0, -1):  # reverse, skip header
        cells = [c.text.strip() for c in rows[i].cells]
        order_date = cells[0] if cells else ""
        if not order_date:  # skip empty rows
            continue
        if _in_range(order_date):
            tr = rows[i]._tr
            tr.getparent().remove(tr)
            deleted += 1

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("clear_rows_from_docx: deleted=%d (range=%s → %s)", deleted, start_date or "*", end_date or "*")
    return buf.read(), deleted


def create_supplier_docx(supplier: dict) -> bytes:
    """Create a new Word document with an order table for a supplier.

    Base columns: Date, Item No., QTY, Color, Customer Name,
                  Sent to Supplier, Ship by date, Sent to customer
    Then one extra column per non-empty custom field name.
    """
    import io as _io

    BASE_COLUMNS = [
        "Date",
        "Item No.",
        "QTY",
        "Color",
        "Customer Name",
        "Sent to Supplier",
        "Ship by date",
        "Sent to customer",
    ]
    custom_cols = [
        cf.get("field_name", "").strip()
        for cf in supplier.get("custom_fields", [])
        if isinstance(cf, dict) and cf.get("field_name", "").strip()
    ]
    all_columns = BASE_COLUMNS + custom_cols

    doc = Document()
    table = doc.add_table(rows=1, cols=len(all_columns))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, col_name in enumerate(all_columns):
        header_cells[i].text = col_name
        para = header_cells[i].paragraphs[0]
        if para.runs:
            para.runs[0].bold = True

    buf = _io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def build_generated_timestamp() -> str:
    """Build UTC timestamp string for diagnostics and logging."""
    return datetime.now(timezone.utc).isoformat()
